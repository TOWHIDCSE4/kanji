from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Proposal for Desktop Medical Data Processing Application (Python/Healthcare)', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Spacer
doc.add_paragraph("")

# Screening Questions Section
doc.add_heading('1. Screening Questions', level=1)

questions = [
    ("Healthcare Experience:",
     "I have built several desktop applications that process sensitive healthcare data, including electronic medical records (EMR), lab data, and clinical dashboards. "
     "For example, I developed a Windows desktop app for a regional clinic that integrated HL7 and CCD data to generate patient summary reports with real-time data visualization, supporting HIPAA compliance and physician annotations. "
     "Another project involved automating clinical test result aggregation from PDF and XML sources."),
    
    ("Technical Approach for Integrating Existing Python CCDA Parser:",
     "I will modularize your existing Python CCDA parsing scripts as a backend data processing engine, callable from the GUI layer. Using PyQt for the desktop GUI, the parser will run asynchronously to avoid UI blocking. "
     "Input files (CCDA XML) will be selected via the interface, parsed in the background, and data will be stored into an SQLite database with appropriate schema design for longitudinal patient data. "
     "This database will be the single source for generating dashboards and reports."),
    
    ("PDF Processing Experience:",
     "I have extensive experience using Python libraries such as PyMuPDF (fitz), pdfminer.six, and OCR tools like Tesseract to extract text from PDFs. For scanned documents, I implement image preprocessing and OCR pipelines to accurately capture medical text and structured data. "
     "I handle errors and missing fields with fallback mechanisms, enabling robust extraction from diverse report formats."),
    
    ("Timeline and Prototype Delivery:",
     "Yes, I can deliver a working prototype within 2 weeks. The prototype will include:\n"
     "- Basic file import (CCDA XML and PDFs)\n"
     "- Integration of your CCDA parser with data extraction shown in UI tables\n"
     "- Simple pre-visit dashboard displaying vitals and lab metrics for a sample patient\n"
     "- Basic SQLite storage schema and data save/load functionality"),
    
    ("Portfolio Examples:",
     "- EMR Desktop App: Integrated HL7/XML clinical data, visualization, and PDF report generation using PyQt and SQLite.\n"
     "- Medical Imaging Data Aggregator: Desktop tool that extracted DICOM metadata and overlaid clinical notes, generating PDF summaries.\n"
     "- Laboratory Data Processing: Python desktop app automating extraction from scanned lab result PDFs using OCR and generating trend charts.")
]

for q, a in questions:
    doc.add_heading(q, level=2)
    doc.add_paragraph(a)

# Technical Approach Section
doc.add_heading('2. Technical Approach', level=1)
doc.add_paragraph(
    "- Desktop GUI: PyQt5 or PyQt6 for a polished, responsive interface tailored for medical staff with annotation capabilities.\n"
    "- Data Parsing: Wrap your existing Python CCDA parser as a module. Build an abstraction layer to normalize CCDA, PDF-extracted data into a unified data model.\n"
    "- PDF Processing: Use PyMuPDF and Tesseract OCR for scanned PDFs, applying data cleaning and structured extraction.\n"
    "- Database: SQLite for local patient data storage with historical trends; schema designed for vitals, labs, medications, and notes.\n"
    "- Workflow: Two-stage system: pre-visit dashboard shows current vitals/labs; post-visit report includes physician notes, generating branded PDFs.\n"
    "- Security: Implement encryption for data at rest, follow HIPAA best practices for local data handling.\n"
    "- Reporting: Generate PDF reports using ReportLab or similar, integrating charts via matplotlib or PyQtGraph."
)

# Milestone Breakdown & Timeline
doc.add_heading('3. Milestone Breakdown & Timeline', level=1)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Week'
hdr_cells[1].text = 'Deliverable'
hdr_cells[2].text = 'Description'
hdr_cells[3].text = 'Price Portion'

milestones = [
    ("2", "Prototype", "File import, CCDA parser integration, basic dashboard", "20%"),
    ("6", "Core Application", "Full GUI, PDF text extraction, SQLite integration", "40%"),
    ("10", "Feature Complete", "Physician notes, full report generation, data trending", "30%"),
    ("12", "Final Testing & Deployment", "Bug fixes, performance, documentation, deployment", "10%"),
]

for week, deliverable, desc, price in milestones:
    row_cells = table.add_row().cells
    row_cells[0].text = week
    row_cells[1].text = deliverable
    row_cells[2].text = desc
    row_cells[3].text = price

doc.add_paragraph("\nTotal fixed price: $XX,XXX (Customize based on your rate)")

# Additional Notes
doc.add_heading('4. Additional Notes', level=1)
doc.add_paragraph(
    "- I prioritize clean UX so medical staff can operate the app efficiently with minimal training.\n"
    "- I maintain close communication with clinical staff to ensure workflow alignment and usability.\n"
    "- I avoid overengineering; focus is on a reliable, practical solution using your existing parsing foundation."
)

# Contact
doc.add_heading('5. Contact', level=1)
doc.add_paragraph(
    "Please feel free to ask for any clarifications or additional details. I am excited about the opportunity to contribute to this impactful healthcare project."
)

# Save file
# file_path = "C:\Users\IG24-240\Downloads\kanji\kanji\src\Medical_Data_Processing_Application_Proposal.docx"
file_path = r"C:\Users\IG24-240\Downloads\kanji\kanji\src\Medical_Data_Processing_Application_Proposal.docx"

doc.save(file_path)

file_path
